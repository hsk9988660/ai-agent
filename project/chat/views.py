from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import KnowledgeBase
from django.contrib.auth import authenticate
from docx import Document
from rest_framework.authentication import TokenAuthentication
from rest_framework.permissions import IsAuthenticated
import logging
import re
import openai
from dotenv import load_dotenv
import os
load_dotenv()

openai.api_key  = os.getenv("OPENAI_API_KEY")

    
class KnowledgeBaseUploadView(APIView):
    # authentication_classes = [TokenAuthentication]
    # permission_classes = [IsAuthenticated]  # Ensure the user is authenticated

    def post(self, request):
        logging.warning(f"Request user: {request.user}")

        # Ensure the user is an admin
        print(f"Request user: {request.user}")  # Add this for debugging

        # if not request.user.is_staff:
        #     return Response({"error": "Only admin users can upload knowledge base files."}, status=status.HTTP_403_FORBIDDEN)

        file = request.FILES.get('file')
        if not file:
            return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)

        content = self.extract_content(file)
        if content:
            # Replace existing knowledge base file
            KnowledgeBase.objects.all().delete()
            knowledge_base = KnowledgeBase.objects.create(content=content)
            return Response({"message": "Knowledge base uploaded successfully!", "file_id": knowledge_base.id}, status=status.HTTP_201_CREATED)
        
        return Response({"error": "Invalid file type"}, status=status.HTTP_400_BAD_REQUEST)

    def extract_content(self, file):
        if file.name.endswith('.txt'):
            return file.read().decode('utf-8')
        elif file.name.endswith('.docx'):
            doc = Document(file)
            return " ".join([p.text for p in doc.paragraphs])
        return None

class KnowledgeBaseFileView(APIView):
    def get(self, request):
        """
        Fetches all knowledge base files.
        """
        knowledge_base = KnowledgeBase.objects.all()
        if not knowledge_base:
            return Response({"error": "No knowledge base files found."}, status=status.HTTP_404_NOT_FOUND)

        # Return all files with their ID and preview
        files = [{"id": kb.id, "content_preview": kb.content[:100]} for kb in knowledge_base]
        return Response({"files": files}, status=status.HTTP_200_OK)

    def delete(self, request, pk=None):
        """
        Deletes a specific knowledge base file by ID.
        """
        try:
            knowledge_base = KnowledgeBase.objects.get(pk=pk)
            knowledge_base.delete()
            return Response({"message": f"Knowledge base file with ID {pk} deleted successfully."}, status=status.HTTP_200_OK)
        except KnowledgeBase.DoesNotExist:
            return Response({"error": f"Knowledge base file with ID {pk} not found."}, status=status.HTTP_404_NOT_FOUND)

    def put(self, request, pk=None):
        """
        Updates a specific knowledge base file by ID.
        """
        if not pk:
            return Response({"error": "File ID must be provided for update."}, status=status.HTTP_400_BAD_REQUEST)

        file = request.FILES.get('file')
        if not file:
            return Response({"error": "No file provided."}, status=status.HTTP_400_BAD_REQUEST)

        content = self.extract_content(file)
        if not content:
            return Response({"error": "Invalid file type."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            knowledge_base = KnowledgeBase.objects.get(pk=pk)
            knowledge_base.content = content
            knowledge_base.save()
            return Response({"message": f"Knowledge base file with ID {pk} updated successfully."}, status=status.HTTP_200_OK)
        except KnowledgeBase.DoesNotExist:
            return Response({"error": f"Knowledge base file with ID {pk} not found."}, status=status.HTTP_404_NOT_FOUND)

    def extract_content(self, file):
        """
        Extracts content from uploaded .txt or .docx files.
        """
        if file.name.endswith('.txt'):
            return file.read().decode('utf-8')  # Handle text files
        elif file.name.endswith('.docx'):
            doc = Document(file)
            return " ".join([p.text for p in doc.paragraphs])  # Handle Word documents
        return None


class AdminLoginView(APIView):
    """
    Admin login endpoint. Authenticates admin users based on username and password.
    """
    def post(self, request):
        # Extract username and password from the request data
        username = request.data.get('username')
        password = request.data.get('password')

        # Authenticate the user
        user = authenticate(username=username, password=password)

        if user:
            if user.is_staff:  # Ensure the user has admin privileges
                return Response(
                    {"message": "Login successful"},
                    status=status.HTTP_200_OK
                )
            return Response(
                {"error": "You do not have admin privileges"},
                status=status.HTTP_403_FORBIDDEN
            )
        return Response(
            {"error": "Invalid credentials"},
            status=status.HTTP_401_UNAUTHORIZED
        )
    
def preprocess_knowledge_base(knowledge_entries):
    """
    Preprocess and split knowledge base content into clean paragraphs.
    """
    paragraphs = []
    for entry in knowledge_entries:
        content = entry.strip()
        # Split paragraphs using double newlines or full stops
        split_content = re.split(r"\n\n|\.\s+", content)
        cleaned_paragraphs = [para.strip() for para in split_content if len(para.strip()) > 20]  # Min length = 20
        paragraphs.extend(cleaned_paragraphs)
    return paragraphs


class QueryView(APIView):
    """
    Handles user queries using ChatGPT (GPT-4 or GPT-3.5-turbo).
    """

    def post(self, request):
        query = request.data.get("query", "").strip()
        if not query:
            return Response({"error": "No query provided."}, status=status.HTTP_400_BAD_REQUEST)

        # Fetch and prepare the knowledge base
        knowledge_base = self.get_knowledge_base()
        if not knowledge_base:
            return Response(
                {"response": "The knowledge base is empty. Please contact the admin."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Combine query with the knowledge base
        combined_context = self.prepare_combined_context(knowledge_base)

        # Use ChatGPT to generate an answer
        response = self.answer_with_chatgpt(query, combined_context)

        return Response({"response": response}, status=status.HTTP_200_OK)

    def get_knowledge_base(self):
        """
        Fetch knowledge base entries and return as a combined string.
        """
        try:
            knowledge_entries = KnowledgeBase.objects.all()
            paragraphs = [entry.content.strip() for entry in knowledge_entries if entry.content.strip()]
            return "\n\n".join(paragraphs) if paragraphs else None
        except Exception as e:
            logging.error(f"Error fetching knowledge base: {e}")
            return None

    def prepare_combined_context(self, knowledge_base):
        """
        Format the knowledge base for better input to ChatGPT.
        """
        return (
            "You are a knowledgeable assistant. Use the following knowledge base to answer questions:\n\n"
            f"{knowledge_base}\n\n"
            "Please respond to the user's query clearly and concisely using the above information."
        )

    def answer_with_chatgpt(self, query, combined_context):
        """
        Use ChatGPT (GPT-4 or GPT-3.5-turbo) to generate a response.
        """
        try:
            # Prepare messages for the OpenAI ChatGPT API
            messages = [
                {"role": "system", "content": combined_context},
                {"role": "user", "content": query}
            ]

            # Call the ChatGPT model
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",  # Use "gpt-3.5-turbo" if GPT-4 is unavailable
                messages=messages,
                max_tokens=500,
                temperature=0.7
            )

            # Extract the content from the response
            return response['choices'][0]['message']['content'].strip()
        except Exception as e:
            logging.error(f"Error generating response from ChatGPT: {e}")
            return "Sorry, I couldn't process your query at the moment."
            
        
class AdminLogoutView(APIView):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        request.auth.delete()  # Deletes the token
        return Response({"message": "Logged out successfully!"}, status=status.HTTP_200_OK)
